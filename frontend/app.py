# frontend/app.py (unchanged from previous, but ensure sys.path is correct)
import sys
import os
from flask import Flask, render_template, request, session, send_file
from flask_bootstrap import Bootstrap5
import pandas as pd
import io
import time
import hashlib
from io import StringIO
from dotenv import load_dotenv
from werkzeug.utils import secure_filename
import base64
import json
import re
from typing import List, Tuple, Optional, Dict
import sqlparse
from concurrent.futures import ThreadPoolExecutor, as_completed
from docx import Document
from docx.shared import Inches
import tempfile

load_dotenv()

app = Flask(__name__)
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'dev_key')
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(__file__), 'static', 'uploads')
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
bootstrap = Bootstrap5(app)

# In-memory cache to avoid storing large payloads in client-side session cookies
# Keys are file_hash or table_name; values are the full autonomous report sections
REPORT_CACHE = {}

# Fix path: since app.py in frontend/, backend is ../backend
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

from backend.data_ingestion.excel_loader import load_excel_to_mysql, drop_table
from backend.query_engine.parser import natural_to_sql
from backend.query_engine.executor import (
    execute_query_mysql, prepare_query, generate_visualization,
    explain_output, execute_advanced_analysis
)
from backend.query_engine.insights import basic_describe
from backend.query_engine.memory import memory
from backend.llm.ollama_client import query_gemini
from backend.utils.logger import get_logger

logger = get_logger(__name__)

MAX_FILE_SIZE_MB = int(os.getenv("MAX_FILE_SIZE_MB", 50))
ALLOWED_EXTENSIONS = {'xlsx', 'xls', 'csv'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def df_to_excel(df):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)
    output.seek(0)
    return output.getvalue()

def encode_image(buf):
    try:
        buf.seek(0)
        data = buf.read()
        logger.debug(f"encode_image: read {len(data)} bytes from buffer")
        if len(data) == 0:
            logger.error("encode_image: buffer data is empty!")
            return None
        b64 = base64.b64encode(data).decode('utf-8')
        logger.debug(f"encode_image: encoded to base64, length {len(b64)}")
        return b64
    except Exception as e:
        logger.error(f"encode_image failed: {e}", exc_info=True)
        return None

def safe_read_json(json_str, orient='split'):
    if json_str:
        return pd.read_json(StringIO(json_str), orient=orient)
    return pd.DataFrame()

def clean_sql(sql_str: str) -> str:
    # Clean up LLM output - extract pure SQL
    sql_str = re.sub(r"```(?:sql)?", "", sql_str, flags=re.IGNORECASE).strip()
    
    # Remove any explanatory text before SELECT
    sql_str = re.sub(r"^.*?(?=SELECT)", "", sql_str, flags=re.IGNORECASE | re.DOTALL).strip()
    
    # Extract just the SQL statement (if LLM adds extra text)
    match = re.search(r"(SELECT .*?;)", sql_str, flags=re.IGNORECASE | re.DOTALL)
    sql = match.group(1).strip() if match else sql_str
    
    # Ensure it starts with SELECT and ends with semicolon
    if not sql.upper().startswith('SELECT'):
        sql = re.sub(r"^.*?(SELECT)", r"\1", sql, flags=re.IGNORECASE | re.DOTALL)
    
    if not sql.endswith(';'):
        sql += ';'
    
    return sql

def extract_sql_columns(sql_text: str):
    """
    Safely extract column expressions from a SELECT statement.
    - Handles messy formatting or extra text
    - Returns a list of column strings
    - Returns [] if parsing fails
    """
    if not sql_text or not sql_text.strip():
        logger.error("Empty or None SQL string")
        return []

    cleaned = sql_text.strip()

    # Try to extract SELECT ... FROM portion using regex
    match = re.search(r"SELECT\s+(.*?)\s+FROM", cleaned, re.IGNORECASE | re.DOTALL)
    if match:
        select_fragment = match.group(1)
        try:
            # Use sqlparse to split columns safely
            parsed = sqlparse.parse(f"SELECT {select_fragment}")[0]
            columns = []
            for token in parsed.tokens:
                if token.ttype is None and token.is_group:
                    for t in token.tokens:
                        if t.ttype is None or t.ttype in (sqlparse.tokens.Name, sqlparse.tokens.Keyword):
                            col = t.value.strip()
                            if col and col != ',':
                                columns.append(col)
            return columns
        except Exception as e:
            logger.error(f"Failed to parse SELECT columns: {e}")
            return []

    logger.error("No valid SELECT statement found")
    return []

def should_show_table(user_query: str, df: pd.DataFrame) -> bool:
    """
    Determine if a table should be displayed for query results.
    ALWAYS shows the table if data is available, because:
    1. Tables provide clear data representation
    2. Users may want to see the actual data, not just explanations
    3. Both visualizations AND tables are valuable for analysis
    """
    # If dataframe is empty or None, don't show table
    if df is None or df.empty:
        return False
    
    # ALWAYS show table if data is available (has at least 1 row)
    # This ensures visualizations from LLM queries are displayed as tables
    return True

def safe_json_loads(text: str):
    """
    Safely load a JSON string from potentially messy LLM output.
    - Strips whitespace
    - Attempts to extract JSON array/object with regex if extra text is included
    - Returns [] if parsing fails
    """
    if not text or not text.strip():
        logger.error("Empty or None JSON string")
        return []

    cleaned = text.strip()

    # Try to extract the JSON portion from messy text (like “Here’s your JSON: [...]”)
    match = re.search(r'(\{.*\}|\[.*\])', cleaned, re.DOTALL)
    if match:
        json_fragment = match.group(1)
        try:
            return json.loads(json_fragment)
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse extracted JSON fragment: {e}")
            return []

    # Still nothing valid
    logger.error("No valid JSON found in text")
    return []

def format_explanation_to_html(text: str) -> str:
    """
    Convert a plain-text, bullet-style explanation into simple HTML:
    - Bold markers **...** become <strong>...</strong>
    - Lines starting with bullet symbols (•, -, *) become <li> items
    - Groups bullet items into a single <ul>
    - Other lines become separate <p> blocks
    This is a lightweight formatter to ensure point-wise, line-broken output
    renders nicely without introducing a markdown dependency.
    """
    if not text:
        return ""
    try:
        import re
        # Normalize line breaks
        lines = [l.strip() for l in text.strip().splitlines() if l.strip()]
        # Bold **...**
        def boldify(s: str) -> str:
            return re.sub(r"\*\*(.*?)\*\*", r"<strong>\1</strong>", s)

        items_html = []
        list_buffer = []
        def flush_list():
            nonlocal items_html, list_buffer
            if list_buffer:
                items_html.append("<ul>" + "".join(f"<li>{boldify(i)}</li>" for i in list_buffer) + "</ul>")
                list_buffer = []

        for line in lines:
            # Detect bullet lines
            if line.startswith(('• ', '- ', '* ')):
                content = line[2:].strip()
                list_buffer.append(content)
            else:
                # New paragraph, flush any existing list
                flush_list()
                items_html.append(f"<p>{boldify(line)}</p>")
        # Flush trailing list
        flush_list()
        return "".join(items_html)
    except Exception:
        # Fallback: simple <br> join and bold conversion
        try:
            text = re.sub(r"\*\*(.*?)\*\*", r"<strong>\1</strong>", text)
        except Exception:
            pass
        return "<br>".join(text.splitlines())

def sanitize_sql_common_issues(sql: str) -> str:
    """
    Reduce common LLM SQL issues that trigger MySQL 1064 errors.
    - Remove unnecessary parentheses around UNION parts
    - Collapse ") UNION ALL (" into " UNION ALL "
    - Trim stray wrapping parentheses before trailing semicolon
    """
    if not sql:
        return sql
    try:
        import re
        s = sql.strip()
        # Replace ) UNION ALL ( with UNION ALL
        s = re.sub(r"\)\s+UNION\s+ALL\s+\(", " UNION ALL ", s, flags=re.IGNORECASE)
        # If the whole statement is wrapped in a single pair of parentheses, unwrap
        s = re.sub(r"^\(\s*(SELECT[\s\S]*?)\s*\)\s*;?$", r"\1;", s, flags=re.IGNORECASE)
        # Ensure ends with semicolon
        if not s.endswith(';'):
            s += ';'
        return s
    except Exception:
        return sql

def generate_autonomous_report(table_name: str, columns: list, insights: dict, df_sample_json: str) -> List[Dict]:
    sample_df = safe_read_json(df_sample_json)
    sections = []
    schema_info = f"Table: {table_name}\nColumns: {', '.join([f'{c} ({insights.get('dtypes', {}).get(c, 'unknown')})' for c in columns])}\nSample Data:\n{sample_df.head().to_string()}"
    
    output_format = [
  {
    "title": "Descriptive Section Title",
    "purpose": "Detailed description of what this section analyzes and why it's valuable for business decision-making"
  },
]
    # Perfect prompt for LLM to generate comprehensive report sections focused on visualization
    comprehensive_prompt = f"""You are an expert data analyst and business intelligence specialist. Create a comprehensive report structure for this dataset that prioritizes VISUALIZABLE INSIGHTS and CHARTS.

DATASET SCHEMA:
{schema_info}

VISUALIZATION-FOCUSED ANALYSIS REQUIREMENTS:
1. **Prioritize Visualizable Data**: Focus on analyses that create meaningful charts, graphs, and visual insights
2. **Avoid Simple Aggregations**: Skip single-row summaries (like total sales, average profit) - these don't create useful visualizations
3. **Create Report Sections**: Design 6-8 sections that generate VISUAL INSIGHTS:
   - **Top/Bottom Performers**: Rankings and comparisons (bar charts, horizontal bars)
   - **Trend Analysis**: Time-based patterns and growth (line charts, area charts)
   - **Category Breakdowns**: Grouped analysis by segments (pie charts, stacked bars)
   - **Distribution Analysis**: Data spread and patterns (histograms, box plots)
   - **Correlation Analysis**: Relationships between variables (scatter plots)
   - **Geographic Analysis**: Location-based insights (maps, regional comparisons)
   - **Seasonal Patterns**: Time-based cyclical analysis (line charts with trends)
   - **Performance Comparisons**: Before/after, category vs category (comparative charts)

4. **Visualization Requirements**: Each section must produce data suitable for charts:
   - Multiple rows (5-50 rows ideal for visualization)
   - Categorical + numeric combinations
   - Time series data (if available)
   - Ranking/ordering data
   - Comparative data across categories

5. **Avoid These**: Single-row summaries, simple totals, basic averages that don't create visual value

OUTPUT FORMAT: Return ONLY a valid JSON array with this structure:
{output_format}

Generate 6-8 sections that create VISUAL INSIGHTS and CHARTS from this dataset."""
    
    try:
        sections_json = query_gemini(comprehensive_prompt)
        # If LLM responded with quota/429 message, force fallback
        if sections_json and any(k in sections_json.lower() for k in ["quota exceeded", "llm quota", "429"]):
            raise ValueError("LLM quota exceeded")
        sections_list = safe_json_loads(sections_json)
        if not sections_list:
            raise ValueError("Empty sections from LLM")
    except Exception as e:
        logger.error(f"Failed to parse sections JSON: {e}")
        # Visualization-focused fallback
        sections_list = [
            {"title": "Top Performing Categories", "purpose": "Ranking analysis showing the best performing categories, products, or segments with comparative metrics for visualization."},
            {"title": "Performance Trends Over Time", "purpose": "Time-based analysis showing trends, patterns, and seasonal variations in key metrics for trend visualization."},
            {"title": "Category Breakdown Analysis", "purpose": "Detailed breakdown of performance across different categories, segments, or groups for comparative visualization."},
            {"title": "Distribution Analysis", "purpose": "Analysis of data distribution patterns, ranges, and statistical spread for histogram and distribution charts."},
        ]
    
    def process_section(sec: Dict) -> Optional[Dict]:
        title = sec['title']
        purpose = sec['purpose']

        # Perfect SQL generation prompt for VISUALIZABLE insights
        sql_prompt = f"""You are an expert SQL analyst. Generate the perfect MySQL SELECT query for VISUALIZABLE business analysis.

BUSINESS ANALYSIS REQUEST:
- Title: {title}
- Purpose: {purpose}

DATABASE SCHEMA:
- Table: {table_name}
- Available Columns: {', '.join(columns)}

VISUALIZATION-FOCUSED SQL REQUIREMENTS:
1. **Generate VISUALIZABLE Data**: Create queries that produce multiple rows (5-50 rows) suitable for charts
2. **Use GROUP BY**: Always group by categorical columns to create comparative data
3. **Create Rankings**: Use ORDER BY and LIMIT to show top/bottom performers
4. **Time Analysis**: If date columns exist, use DATE_FORMAT() for time-based grouping
5. **Comparative Analysis**: Group by categories to compare performance across segments

ACCURATE SQL PATTERNS (use exact column names from schema):
- **Top Performers**: SELECT [categorical_column], SUM([numeric_column]) FROM {table_name} GROUP BY [categorical_column] ORDER BY SUM([numeric_column]) DESC LIMIT 10
- **Trend Analysis**: SELECT DATE_FORMAT([date_column], '%Y-%m') as period, SUM([numeric_column]) FROM {table_name} GROUP BY period ORDER BY period
- **Category Breakdown**: SELECT [categorical_column], COUNT(*), AVG([numeric_column]) FROM {table_name} GROUP BY [categorical_column] ORDER BY COUNT(*) DESC LIMIT 15
- **Distribution**: SELECT [categorical_column], COUNT(*) FROM {table_name} GROUP BY [categorical_column] ORDER BY COUNT(*) DESC LIMIT 20

CRITICAL ACCURACY REQUIREMENTS:
- MUST use exact column names from the provided schema: {', '.join(columns)}
- MUST use GROUP BY to create multiple rows
- MUST use ORDER BY for meaningful sorting
- MUST use LIMIT to focus results (10-20 rows typically)
- AVOID single-row aggregations (SUM without GROUP BY)
- Focus on categorical + numeric combinations
- Use only existing columns from the schema - NO made-up column names
- Ensure MySQL 8.x compatibility
- Double-check column names match exactly with schema

OUTPUT FORMAT:
- Start with SELECT (not "SELECT query" or any other text)
- End with semicolon
- No explanatory text before or after the SQL
- No comments or descriptions

Generate SQL that creates VISUALIZABLE DATA with multiple rows for charts and graphs."""
        
        try:
            sql = query_gemini(sql_prompt)
            sql = clean_sql(sql)
            sql = sanitize_sql_common_issues(sql)
            if not (sql and 'SELECT' in sql.upper()):
                return None
            try:
                df = execute_query_mysql(sql)
                if df is None or df.empty:
                    logger.warning(f"Section {title}: Query returned empty result, skipping")
                    return None
            except Exception as query_error:
                logger.error(f"Section {title}: Query execution failed: {query_error}")
                return None

            # Check if data is suitable for visualization (more than 1 row)
            is_visualizable = df.shape[0] > 1

            # Concise, point-wise explanation prompt
            explanation_prompt = f"""You are a senior business analyst. Provide concise, point-wise insights for this report section.

SECTION DETAILS:
- Title: {title}
- Purpose: {purpose}
- Data Shape: {df.shape[0]} rows, {df.shape[1]} columns
- Columns: {list(df.columns)}

SAMPLE DATA:
{df.head(5).to_string()}

REQUIREMENTS:
1. **Format**: Each point on a new line with bullet points
2. **Length**: Keep each point brief but descriptive
3. **Content**: Focus on key insights, patterns, and actionable recommendations
4. **Numbers**: Include specific data points and metrics
5. **Structure**: Use clear bullet points, not paragraphs

FORMAT EXAMPLE:
• Key finding 1 with specific numbers
• Key finding 2 with specific numbers  
• Business implication 1
• Business implication 2
• Recommendation 1
• Recommendation 2

Provide concise, point-wise insights with specific numbers. Each point should be brief but informative."""
                    
            try:
                explanation_raw = query_gemini(explanation_prompt)
                if not explanation_raw or explanation_raw.strip() == "":
                    explanation_raw = f"**{title}:**\n• Analysis completed.\n• Rows: {df.shape[0]}\n• Columns: {df.shape[1]}"
            except Exception as exp_error:
                logger.warning(f"Explanation generation failed for {title}: {exp_error}")
                explanation_raw = f"**{title}:**\n• Analysis completed.\n• Rows: {df.shape[0]}\n• Columns: {df.shape[1]}"
            # Ensure point-wise HTML formatting
            explanation = format_explanation_to_html(explanation_raw)

            # Generate visualization for all data
            viz_b64 = None
            chart_b64s = []
            
            # Always try to generate visualization if we have data
            if df is not None and not df.empty:
                try:
                    viz_buf = generate_visualization(df)
                    if viz_buf:
                        try:
                            buf_data = viz_buf.getvalue()
                            if buf_data and len(buf_data) > 0:
                                viz_b64 = encode_image(viz_buf)
                                logger.info(f"Section {title}: Generated visualization ({len(buf_data)} bytes)")
                            else:
                                logger.warning(f"Section {title}: Visualization buffer empty")
                        except Exception as e:
                            logger.warning(f"Section {title}: Failed to encode visualization: {e}")
                except Exception as viz_error:
                    logger.warning(f"Section {title}: Visualization generation failed: {viz_error}", exc_info=False)
                    viz_b64 = None
                
                # Try to generate additional charts from advanced analysis
                try:
                    _, charts = execute_advanced_analysis(df, purpose, generate_charts=True)
                    chart_b64s = []
                    for t, b in charts:
                        try:
                            if b:
                                buf_data = b.getvalue()
                                if buf_data and len(buf_data) > 0:
                                    chart_b64s.append((t, encode_image(b)))
                        except Exception as chart_err:
                            logger.warning(f"Section {title}: Failed to encode chart {t}: {chart_err}")
                except Exception as advanced_error:
                    logger.warning(f"Section {title}: Advanced analysis failed: {advanced_error}")
                    chart_b64s = []
            else:
                logger.warning(f"Section {title}: No data to visualize")

            # Create properly styled table HTML with responsive wrapper
            table_html = df.to_html(classes='table table-striped table-hover table-sm', escape=False, index=False)
            # Wrap in responsive container with better styling
            df_html = f"""
            <div class="table-responsive" style="max-height: 400px; overflow-y: auto;">
                {table_html}
            </div>
            """

            return {
                'title': title,
                'explanation': explanation,
                'df_html': df_html,
                'viz_b64': viz_b64,
                'chart_b64s': chart_b64s
            }
        except Exception as e:
            logger.error(f"Section {title} failed: {e}", exc_info=True)
            return None

    # Parallel execution of sections
    max_workers = int(os.getenv('REPORT_MAX_WORKERS', '4'))
    futures = []
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        for sec in sections_list[:6]:
            futures.append(executor.submit(process_section, sec))
        for fut in as_completed(futures, timeout=180):  # 3 minute timeout per section
            try:
                res = fut.result(timeout=180)
                if res:
                    sections.append(res)
            except Exception as e:
                logger.error(f"Parallel section failed: {e}")
    
    return sections


def _add_html_paragraphs_to_doc(doc: Document, html_text: str):
    """
    Lightweight converter: convert simple formatted explanation (HTML with <p>, <ul>, <li>, <strong>)
    into Word paragraphs and bold runs.
    """
    if not html_text:
        return
    # Very small HTML handling without adding a heavy dependency
    try:
        # Replace common tags with simple markers
        text = html_text.replace('<strong>', '**').replace('</strong>', '**')
        # Handle unordered lists
        text = text.replace('<ul>', '\n').replace('</ul>', '\n')
        text = text.replace('<li>', '• ').replace('</li>', '\n')
        # Remove remaining tags like <p>
        text = re.sub(r'<[^>]+>', '', text)
        # Split into lines and create paragraphs
        for line in [l.strip() for l in text.splitlines() if l.strip()]:
            # Handle bold markers
            parts = re.split(r'(\*\*.*?\*\*)', line)
            p = doc.add_paragraph()
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
    except Exception as e:
        # Fallback: add raw text
        doc.add_paragraph(re.sub(r'<[^>]+>', '', html_text))


def create_report_docx(report_sections: List[Dict]) -> bytes:
    """
    Build a .docx binary from report sections structure.
    Each section is expected to contain: title, explanation (HTML), df_html (optional), viz_b64 (optional), chart_b64s (optional list of (title, b64)).
    """
    doc = Document()
    doc.add_heading('Autonomous Report', level=1)
    for sec in report_sections:
        title = sec.get('title', 'Section')
        doc.add_heading(title, level=2)
        explanation_html = sec.get('explanation', '')
        _add_html_paragraphs_to_doc(doc, explanation_html)

        # Add table HTML if present - convert the pandas-generated HTML table into a native docx table
        df_html = sec.get('df_html')
        if df_html:
            try:
                # Find the first <table>...</table> block
                table_match = re.search(r'<table.*?>(.*?)</table>', df_html, flags=re.DOTALL | re.IGNORECASE)
                if table_match:
                    table_html = table_match.group(1)
                    # Try to extract header cells
                    headers = re.findall(r'<th[^>]*>(.*?)</th>', table_html, flags=re.DOTALL | re.IGNORECASE)
                    # All rows (including header row(s))
                    all_rows = re.findall(r'<tr>(.*?)</tr>', table_html, flags=re.DOTALL | re.IGNORECASE)

                    if headers:
                        cols = len(headers)
                        table = doc.add_table(rows=1, cols=cols)
                        hdr_cells = table.rows[0].cells
                        for i, h in enumerate(headers):
                            text = re.sub(r'<[^>]+>', '', h).strip()
                            hdr_cells[i].text = text
                            # Bold the header run
                            for run in hdr_cells[i].paragraphs[0].runs:
                                run.bold = True

                        # Add body rows (look for <td> within each <tr>)
                        for r_html in all_rows:
                            # skip header rows that use <th>
                            if re.search(r'<th', r_html, flags=re.IGNORECASE):
                                continue
                            cells = re.findall(r'<td[^>]*>(.*?)</td>', r_html, flags=re.DOTALL | re.IGNORECASE)
                            cells = [re.sub(r'<[^>]+>', '', c).strip() for c in cells]
                            if cells:
                                row_cells = table.add_row().cells
                                for i, c in enumerate(cells[:cols]):
                                    row_cells[i].text = c
                    else:
                        # No header row; build table from all rows
                        all_rows = re.findall(r'<tr>(.*?)</tr>', table_html, flags=re.DOTALL | re.IGNORECASE)
                        # Determine number of columns from first row
                        if all_rows:
                            first_cells = re.findall(r'<t[dh][^>]*>(.*?)</t[dh]>', all_rows[0], flags=re.DOTALL | re.IGNORECASE)
                            cols = len(first_cells) if first_cells else 1
                            table = doc.add_table(rows=0, cols=cols)
                            for r_html in all_rows:
                                cells = re.findall(r'<t[dh][^>]*>(.*?)</t[dh]>', r_html, flags=re.DOTALL | re.IGNORECASE)
                                cells = [re.sub(r'<[^>]+>', '', c).strip() for c in cells]
                                if cells:
                                    row_cells = table.add_row().cells
                                    for i, c in enumerate(cells[:cols]):
                                        row_cells[i].text = c
            except Exception:
                pass

        # Add main visualization image if present
        viz_b64 = sec.get('viz_b64')
        if viz_b64:
            try:
                img_data = base64.b64decode(viz_b64)
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                    tmp.write(img_data)
                    tmp.flush()
                    doc.add_picture(tmp.name, width=Inches(6))
            except Exception:
                pass

        # Additional charts
        for chart in sec.get('chart_b64s', []):
            try:
                # chart may be tuple (title, b64)
                if isinstance(chart, (list, tuple)) and len(chart) == 2:
                    ctitle, cb64 = chart
                else:
                    ctitle, cb64 = ('Chart', chart)
                if ctitle:
                    doc.add_paragraph(ctitle, style='Intense Quote')
                img_data = base64.b64decode(cb64)
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                    tmp.write(img_data)
                    tmp.flush()
                    doc.add_picture(tmp.name, width=Inches(6))
            except Exception:
                continue

    # Write to bytes
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

@app.route('/', methods=['GET', 'POST'])
def index():
    if 'table_name' not in session:
        session['table_name'] = None
        session['columns'] = []
        session['df_sample'] = None
        session['uploaded_files'] = []
        session['file_hash'] = None
        session['auto_insights_generated'] = False
        session['active_tab'] = 'preview'
        session['report_data'] = None
    df_result_global = None
    active_tab = session.get('active_tab', 'preview')
    message = None
    report_data = session.get('report_data', None)
    if request.method == 'POST':
        if 'files' in request.files:
            files = request.files.getlist('files')
            files = [f for f in files if f and allowed_file(f.filename)]
            if files:
                total_size = 0
                for f in files:
                    f.seek(0)
                    total_size += len(f.read())
                    f.seek(0)
                if total_size > MAX_FILE_SIZE_MB * 1024 * 1024:
                    message = "Total size too large."
                else:
                    file_hash = hashlib.md5(str(sorted([f.filename for f in files])).encode()).hexdigest()
                    if session.get('file_hash') != file_hash:
                        unique_suffix = int(time.time())
                        table_name = f"data_fusion_{unique_suffix}"
                        try:
                            table_name, columns, sample_df = load_excel_to_mysql(files, table_name)
                            session.update({
                                'table_name': table_name,
                                'columns': columns,
                                'df_sample': sample_df.to_json(orient='split'),
                                'uploaded_files': [f.filename for f in files],
                                'file_hash': file_hash,
                                'auto_insights_generated': False,
                                'report_data': None
                            })
                            message = f"✅ Loaded {len(files)} files into `{table_name}`."
                            session['active_tab'] = 'preview'
                        except Exception as e:
                            message = f"Load failed: {e}"
                    else:
                        message = "📁 Using cached data."
        elif 'reset_data' in request.form:
            if session.get('table_name'):
                drop_table(session['table_name'])
            # Clear server-side cached report for this dataset
            try:
                cache_key = session.get('file_hash') or session.get('table_name')
                if cache_key and cache_key in REPORT_CACHE:
                    del REPORT_CACHE[cache_key]
            except Exception:
                pass
            session.clear()
            memory.history.clear()
            session['active_tab'] = 'preview'
            message = "🔄 Reset complete."
        elif 'clear_memory' in request.form:
            memory.history.clear()
            message = "🧹 Memory cleared."
        elif 'generate_report' in request.form and session.get('table_name'):
            try:
                insights = session.get('insights', {})
                report_data = generate_autonomous_report(
                    session['table_name'], 
                    session['columns'], 
                    insights, 
                    session['df_sample']
                )
                # Cache server-side to avoid huge cookies
                try:
                    cache_key = session.get('file_hash') or session.get('table_name')
                    if cache_key:
                        REPORT_CACHE[cache_key] = report_data
                except Exception:
                    pass
                session['active_tab'] = 'report'
                message = f"🤖 Generated {len(report_data)} autonomous report sections."
            except Exception as e:
                # Suppress raw error details from UI; log internally
                logger.error(f"Report generation failed: {e}")
                message = "Report generation encountered issues. Please try again."
                session['active_tab'] = 'report'
        elif 'query' in request.form and session.get('table_name'):
            user_query = request.form['query'].strip()
            if user_query:
                if any(phrase in user_query.lower() for phrase in ['full report', 'comprehensive analysis', 'autonomous report', 'generate report']):
                    # Redirect to generate_report logic
                    try:
                        insights = session.get('insights', {})
                        report_data = generate_autonomous_report(
                            session['table_name'], 
                            session['columns'], 
                            insights, 
                            session['df_sample']
                        )
                        # Cache server-side to avoid huge cookies
                        try:
                            cache_key = session.get('file_hash') or session.get('table_name')
                            if cache_key:
                                REPORT_CACHE[cache_key] = report_data
                        except Exception:
                            pass
                        session['active_tab'] = 'report'
                        message = f"🤖 Generated {len(report_data)} autonomous report sections."
                        return render_template(
                            'index.html',
                            table_name=session['table_name'],
                            df_sample=safe_read_json(session['df_sample'], 'split'),
                            insights=insights,
                            message=message,
                            report_data=report_data,
                            active_tab='report'
                        )
                    except Exception as e:
                        logger.error(f"Report generation failed: {e}")
                        message = "Report generation encountered issues. Please try again."
                else:
                    try:
                        sql = natural_to_sql(user_query, session['table_name'], session['columns'])
                        if sql:
                            sql = prepare_query(sql, session['uploaded_files'])
                            # Cache-first: return previous deterministic result for same query + table + file set
                            cached = memory.get_cached(user_query, session.get('table_name') or '', session.get('file_hash') or '')
                            if cached is not None:
                                df_result_global = pd.DataFrame(cached)
                            else:
                                df_result_global = execute_query_mysql(sql)
                            try:
                                session['df_result'] = df_result_global.to_json(orient='split')
                            except Exception:
                                session['df_result'] = None
                            # Conditionally generate visualization only when requested
                            viz_b64 = None
                            generate_viz = ('generate_viz' in request.form and request.form.get('generate_viz') == '1')
                            
                            # Always try to generate visualization if we have data
                            viz_b64 = None
                            if not df_result_global.empty and len(df_result_global) >= 2:
                                logger.info(f"Attempting visualization for {len(df_result_global)} rows...")
                                try:
                                    viz_buf = generate_visualization(df_result_global)
                                    if viz_buf:
                                        buf_data = viz_buf.getvalue()
                                        buf_size = len(buf_data)
                                        logger.info(f"✓ Visualization buffer received: {buf_size} bytes")
                                        if buf_size > 100:  # At least 100 bytes for a valid PNG
                                            viz_b64 = encode_image(viz_buf)
                                            if viz_b64:
                                                logger.info(f"✓ Visualization encoded to base64: {len(viz_b64)} chars")
                                            else:
                                                logger.error("encode_image returned None")
                                        else:
                                            logger.warning(f"Visualization buffer too small: {buf_size} bytes (need >100)")
                                            viz_b64 = None
                                    else:
                                        logger.warning("Visualization buffer is None")
                                        viz_b64 = None
                                except Exception as viz_err:
                                    logger.error(f"Auto visualization failed: {viz_err}", exc_info=True)
                                    viz_b64 = None
                            else:
                                logger.info(f"Skipping visualization: empty={df_result_global.empty}, rows={len(df_result_global)}")
                            
                            explanation = explain_output(df_result_global, user_query)
                            # If forecast requested but result is empty, try using full table for forecasting
                            wants_forecast = any(k in user_query.lower() for k in ['forecast', 'predict'])
                            base_df_for_analysis = df_result_global
                            if wants_forecast and df_result_global.empty and session.get('table_name'):
                                try:
                                    base_df_for_analysis = execute_query_mysql(f"SELECT * FROM {session['table_name']}")
                                except Exception:
                                    base_df_for_analysis = df_result_global
                            
                            # Always generate charts when data is available
                            df_forecast, charts = execute_advanced_analysis(base_df_for_analysis, user_query, generate_charts=True)
                            # If we produced a forecast table, use it as the primary result to show and download
                            if wants_forecast and df_forecast is not None and not df_forecast.empty:
                                df_result_global = df_forecast
                                try:
                                    session['df_result'] = df_result_global.to_json(orient='split')
                                except Exception:
                                    pass
                            chart_b64s = [(title, encode_image(buf)) for title, buf in charts]
                            memory.add(user_query, df_result_global.to_dict('records'), session.get('table_name') or '', session.get('file_hash') or '')
                            df_sample = safe_read_json(session['df_sample'], 'split') if session.get('df_sample') else pd.DataFrame()
                            insights = session.get('insights', {})
                            session['active_tab'] = 'query'
                            
                            # Format result table with responsive container
                            result_table_html = None
                            if should_show_table(user_query, df_result_global):
                                table_html = df_result_global.to_html(classes='table table-striped table-hover table-sm', escape=False, index=False)
                                result_table_html = f"""
                                <div class="table-responsive" style="max-height: 500px; overflow-y: auto;">
                                    {table_html}
                                </div>
                                """
                            
                            return render_template(
                                'index.html',
                                table_name=session['table_name'],
                                df_sample=df_sample,
                                insights=insights,
                                message=message,
                                user_query=user_query,
                                df_result=result_table_html,
                                explanation=explanation,
                                viz_b64=viz_b64,
                                chart_b64s=chart_b64s,
                                active_tab='query')
                    except Exception as e:
                        message = f"Query failed: {e}"
                        session['active_tab'] = 'query'
        elif 'download_report' in request.form:
            # Return server-side cached report as a Word .docx file
            try:
                cache_key = session.get('file_hash') or session.get('table_name')
                report_sections = None
                if cache_key:
                    report_sections = REPORT_CACHE.get(cache_key)
                if not report_sections:
                    # Fall back to any session-stored report data
                    report_sections = session.get('report_data')
                if not report_sections:
                    message = "No report available to download."
                else:
                    docx_bytes = create_report_docx(report_sections)
                    return send_file(
                        io.BytesIO(docx_bytes),
                        as_attachment=True,
                        download_name="autonomous_report.docx",
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )
            except Exception as e:
                logger.error(f"Failed to create DOCX: {e}")
                message = "Failed to create Word document."

        elif 'download_sample' in request.form and session.get('df_sample'):
            df = safe_read_json(session['df_sample'], 'split')
            return send_file(
                io.BytesIO(df_to_excel(df)),
                as_attachment=True,
                download_name="data_sample.xlsx",
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
        elif 'download_result' in request.form and session.get('df_result'):
            try:
                df_cached = safe_read_json(session['df_result'], 'split')
            except Exception:
                df_cached = pd.DataFrame()
            return send_file(
                io.BytesIO(df_to_excel(df_cached)),
                as_attachment=True,
                download_name="query_result.xlsx",
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
    # Auto insights
    if session.get('table_name') and not session.get('auto_insights_generated'):
        try:
            full_df = execute_query_mysql(f"SELECT * FROM {session['table_name']}")
            insights = basic_describe(full_df)
            session['insights'] = {
                'shape': insights['shape'],
                'dtypes': {k: str(v) for k, v in insights['dtypes'].items()},
                'sample_dtypes': dict(list(insights['dtypes'].items())[:5])
            }
            session['auto_insights_generated'] = True
        except Exception as e:
            session['insights'] = {'error': str(e)}
    df_sample = safe_read_json(session['df_sample'], 'split') if session.get('df_sample') else pd.DataFrame()
    insights = session.get('insights', {})
    # Pull report from server-side cache to keep cookies small
    try:
        cache_key = session.get('file_hash') or session.get('table_name')
        report_data = REPORT_CACHE.get(cache_key)
    except Exception:
        report_data = None
    return render_template(
        'index.html',
        table_name=session['table_name'],
        df_sample=df_sample,
        insights=insights,
        message=message,
        report_data=report_data,
        active_tab=active_tab
    )

if __name__ == '__main__':
    app.run(debug=True, use_reloader=False, host='0.0.0.0', port=5000)